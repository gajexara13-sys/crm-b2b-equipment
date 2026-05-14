# -*- coding: utf-8 -*-
"""Подвал КП: условия (terms_*) и подпись (signer_*) — переменные docxtpl.

Запуск из каталога backend:
  python tools/patch_kp_van_footer_terms.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


def main() -> None:
    tpl_dir = Path(__file__).resolve().parent.parent / "templates"
    path = next(p for p in tpl_dir.glob("*.docx") if p.name != "kp_docxtpl.docx")
    doc = Document(path)
    cell = doc.tables[1].cell(0, 0)
    joined = "".join(p.text for p in cell.paragraphs)
    if "{{ signer_name }}" in joined:
        print("footer_already_marked:", path)
        return

    # Срок действия цен — новый абзац перед блоком «Страна / Гарантия…»
    cell.paragraphs[12].insert_paragraph_before().text = (
        "{% if terms_price_validity %}{{ terms_price_validity }}{% endif %}"
    )

    footer_lines = [
        "{% if terms_production_country %}Страна: {{ terms_production_country }}{% endif %}",
        "{% if terms_warranty %}Гарантия: {{ terms_warranty }}{% endif %}",
        "{% if terms_delivery %}Условия поставки: {{ terms_delivery }}{% endif %}",
        "{% if terms_lead_time %}Срок поставки: {{ terms_lead_time }}{% endif %}",
        "{% if terms_payment %}Условия оплаты: {{ terms_payment }}{% endif %}",
        "{{ terms_currency_note }}",
        "{% if terms_address_note %}{{ terms_address_note }}{% endif %}",
    ]
    for idx, line in zip(range(13, 20), footer_lines):
        cell.paragraphs[idx].text = line

    cell.paragraphs[20].text = "С уважением,"
    cell.paragraphs[21].text = "{{ signer_name }}"
    cell.paragraphs[22].text = "{{ signer_position }}"
    cell.paragraphs[23].text = "{{ signer_phone }}"

    doc.save(path)
    print("patched:", path)


if __name__ == "__main__":
    main()
