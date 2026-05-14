# -*- coding: utf-8 -*-
"""Шаблон КП: отступ перед подписью + блок подписанта в таблице 1×2.

Запуск из каталога backend:
  python tools/patch_kp_van_signer_table.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _delete_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def main() -> None:
    tpl_dir = Path(__file__).resolve().parent.parent / "templates"
    path = next(p for p in tpl_dir.glob("*.docx") if p.name != "kp_docxtpl.docx")
    doc = Document(path)
    cell = doc.tables[1].cell(0, 0)

    # --- 1. Пустая строка между условиями и «С уважением» ---
    idx_sign = None
    for i, p in enumerate(cell.paragraphs):
        if p.text.strip() == "С уважением,":
            idx_sign = i
            break
    if idx_sign is None:
        raise RuntimeError("Не найден абзац «С уважением,»")

    prev_empty = idx_sign > 0 and not cell.paragraphs[idx_sign - 1].text.strip()
    if not prev_empty:
        cell.paragraphs[idx_sign].insert_paragraph_before()
        idx_sign += 1

    # --- 2. Найти абзацы с переменными подписанта (до вставки таблицы пересканируем)
    def find_signer_paragraphs():
        idxs = []
        for i, p in enumerate(cell.paragraphs):
            t = p.text.strip()
            if "{{ signer_name }}" in t or "{{ signer_position }}" in t or "{{ signer_phone }}" in t:
                idxs.append(i)
        return idxs

    signer_idxs = find_signer_paragraphs()
    if len(signer_idxs) != 3:
        raise RuntimeError(
            f"Ожидались 3 абзаца подписанта, найдено: {len(signer_idxs)}"
        )

    # Таблица в конце ячейки — затем переносим XML перед первым абзацем подписанта
    tbl = cell.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass

    tbl_el = tbl._element
    tbl_el.getparent().remove(tbl_el)

    anchor = cell.paragraphs[signer_idxs[0]]._element
    anchor.addprevious(tbl_el)

    left = tbl.cell(0, 0)
    right = tbl.cell(0, 1)
    # Левая колонка — пустая под оттиск печати (без «М.П.»)
    left.paragraphs[0].text = ""
    left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    right.paragraphs[0].text = "{{ signer_name }}"
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2 = right.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run("{{ signer_position }}")
    p3 = right.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.add_run("{{ signer_phone }}")

    # Удалить старые три абзаца с теми же тегами (с конца — индексы не плывут)
    for i in sorted(find_signer_paragraphs(), reverse=True):
        _delete_paragraph(cell.paragraphs[i])

    # Лишний пустой абзац между «С уважением,» и таблицей подписи
    for i, p in enumerate(cell.paragraphs):
        if p.text.strip() != "С уважением,":
            continue
        if i + 2 < len(cell.paragraphs):
            if not cell.paragraphs[i + 1].text.strip() and not cell.paragraphs[i + 2].text.strip():
                _delete_paragraph(cell.paragraphs[i + 2])
        break

    doc.save(path)
    print("patched:", path)


if __name__ == "__main__":
    main()
