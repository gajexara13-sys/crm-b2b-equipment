# -*- coding: utf-8 -*-
"""Разовая правка шаблона «КП ТОО ВАН №1.docx»: краткий список позиций и цикл описания.

Запуск из каталога backend:
  python tools/patch_kp_van_items_section.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


def _delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def main() -> None:
    tpl_dir = Path(__file__).resolve().parent.parent / "templates"
    path = next(p for p in tpl_dir.glob("*.docx") if p.name != "kp_docxtpl.docx")
    doc = Document(path)
    cell = doc.tables[1].cell(0, 0)
    joined = "".join(p.text for p in cell.paragraphs)

    # --- Краткий перечень позиций после вводного блока ---
    if "{{ items_short_list }}" not in joined:
        cell.paragraphs[6].text = "{{ items_short_list }}"
        for _ in range(7):
            _delete_paragraph(cell.paragraphs[7])

    joined = "".join(p.text for p in cell.paragraphs)
    if "{% for item in items %}" in joined:
        doc.save(path)
        print("already_has_items_loop:", path)
        return

    hi = si = None
    for i, p in enumerate(cell.paragraphs):
        t = p.text
        if hi is None and "ОПИСАНИЕ" in t and "ОБОРУДОВАНИЯ" in t:
            hi = i
        if "СПЕЦИФИКАЦИЯ" in t and "ПОСТАВКУ" in t:
            si = i

    if hi is None or si is None:
        raise RuntimeError(f"Не найдены маркеры разделов (hi={hi}, si={si})")

    tc = cell._tc
    heading_el = cell.paragraphs[hi]._element
    spec_el = cell.paragraphs[si]._element
    kids = list(tc)
    idx_hi = kids.index(heading_el)
    idx_si = kids.index(spec_el)
    for el in kids[idx_hi + 1 : idx_si]:
        tc.remove(el)

    spec_p = cell.paragraphs[hi + 1]
    loop_p = spec_p.insert_paragraph_before()
    loop_p.add_run(
        "{% for item in items %}"
        "{{ item.title }}"
        "{% if item.model %} | Модель: {{ item.model }}{% endif %}"
        " | {{ item.intro }}"
        "{% if item.kit_text %} | Комплект поставки: {{ item.kit_text }}{% endif %}"
        "{% if item.features_block %} | Основные особенности: {{ item.features_block }}{% endif %}"
        "{% if item.specs_block %} | ТЕХНИЧЕСКИЕ ХАРАКТЕРИСТИКИ: {{ item.specs_block }}{% endif %}"
        " | {{ item.figure_caption }} | {{ item.photo }}"
        " ~~~ "
        "{% endfor %}"
    )
    try:
        loop_p.style = "Body Text"
    except Exception:
        pass

    doc.save(path)
    print("patched:", path)


if __name__ == "__main__":
    main()
