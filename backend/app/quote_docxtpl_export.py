"""Экспорт КП через docxtpl (Word + Jinja2).

QUOTE_DOCX_TEMPLATE_PATH — путь к вашему .docx с теми же переменными.
Файл без тегов {% %} не подставит данные: откройте встроенный kp_docxtpl.docx
как образец разметки.

Встроенный шаблон: backend/templates/kp_docxtpl.docx (создаётся автоматически).
При изменении этого модуля удалите старый kp_docxtpl.docx — он пересоздастся.

Переменные шаблона:
  sender_lines          list[str]   — строки реквизитов отправителя
  recipient_lines       list[str]   — строки реквизитов получателя
  quote_date_ru         str         — «20 апреля 2026 г.»
  quote_number          str
  intro / intro_listing / greeting_name — вводный текст и приветствие
  items_short_list      Listing     — краткий перечень позиций («название, модель: …;»)
  items                 list[dict]  — позиции КП (см. ниже)
    item.title, item.model, item.intro, item.kit_text
    item.features_lines  list[str]
    item.specs           list[{param, value}]
    item.figure_caption  str        — «Рис.N Название Модель»
    item.photo           InlineImage | ""  — первое фото позиции
  spec_rows             list[dict]  — строки спецификации
    r.no, r.name, r.price, r.qty, r.sum_full, r.sum_disc
  sum_total_full / sum_total_disc  str
  price_col_title / sum_col_title / disc_col_title  str
  show_discount_column  bool
  terms_*               str         — условия поставки
  signer_name / signer_position / signer_phone  str
"""

from __future__ import annotations

import io
import logging
import os
import urllib.request
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docxtpl import DocxTemplate, InlineImage, Listing

from app.quote_document_export import _currency_note, format_quote_date_ru, money_ru

_log = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────────────────────
# Вспомогательные функции
# ──────────────────────────────────────────────────────────────────────────────

def _normalize_item_export(it: dict) -> dict:
    row = dict(it)
    if not row.get("show_intro"):
        row["intro"] = None
    if not row.get("show_features"):
        row["features_text"] = None
    if not row.get("show_kit"):
        row["kit_text"] = None
    if not row.get("show_specs"):
        row["specs"] = []
    if not row.get("show_photos"):
        row["photo_urls"] = []
    return row


# Расшифровка аббревиатур организационно-правовых форм
_ORG_FORM_EXPAND: dict[str, str] = {
    "ООО":  "Общество с ограниченной ответственностью",
    "АО":   "Акционерное общество",
    "ЗАО":  "Закрытое акционерное общество",
    "ОАО":  "Открытое акционерное общество",
    "ПАО":  "Публичное акционерное общество",
    "НАО":  "Непубличное акционерное общество",
    "ИП":   "Индивидуальный предприниматель",
    "ТОО":  "Товарищество с ограниченной ответственностью",
    "АО (КЗ)": "Акционерное общество",
    "ГУП":  "Государственное унитарное предприятие",
    "МУП":  "Муниципальное унитарное предприятие",
    "НКО":  "Некоммерческая организация",
    "АНО":  "Автономная некоммерческая организация",
}


def _expand_org_form(lf: str) -> str:
    """Расшифровывает аббревиатуру ОПФ, если она известна. Иначе возвращает как есть."""
    return _ORG_FORM_EXPAND.get(lf.strip(), lf)


def _tax_label_for_currency(currency: str | None) -> str:
    c = (currency or "RUB").upper()
    if c in ("KZT", "KZTEN"):
        return "БИН"
    return "ИНН"


def _sender_lines_for_tpl(sender: dict[str, Any], currency: str | None) -> list[str]:
    lines: list[str] = []
    lf = _expand_org_form((sender.get("legal_form") or "").strip())
    ln = (sender.get("legal_name") or sender.get("name") or "").strip()
    if lf:
        lines.append(lf)
    if ln:
        lines.append(f"«{ln}»")
    addr = (sender.get("legal_address") or "").strip()
    if addr:
        lines.append(addr)
    tax_n = (sender.get("tax_number") or "").strip()
    kpp = (sender.get("kpp") or "").strip()
    ogrn = (sender.get("ogrn") or "").strip()
    tl = _tax_label_for_currency(currency)
    tax_parts: list[str] = []
    if tax_n:
        tax_parts.append(f"{tl} {tax_n}")
    if kpp and tl == "ИНН":
        tax_parts.append(f"КПП {kpp}")
    if tax_parts:
        lines.append(", ".join(tax_parts) + ("," if kpp and tl == "ИНН" else ""))
    if ogrn:
        lines.append(f"ОГРН {ogrn}")
    em = (sender.get("email") or "").strip()
    if em:
        lines.append(f"e-mail: {em}")
    ph = (sender.get("phone") or "").strip()
    if ph:
        lines.append(f"Тел.: {ph}")
    web = (sender.get("website") or "").strip()
    if web:
        lines.append(f"Сайт: {web}")
    return lines


def _items_short_listing(items_raw: list[dict[str, Any]]) -> Listing | str:
    """Краткий перечень позиций КП (строки «название, модель: …;») для шапки текста."""
    lines: list[str] = []
    for it in items_raw:
        row = _normalize_item_export(it)
        title = (row.get("title") or "").strip()
        if not title:
            continue
        model = (row.get("model") or "").strip()
        if model:
            lines.append(f"{title}, модель: {model};")
        else:
            lines.append(f"{title};")
    return Listing("\n".join(lines)) if lines else ""


def _recipient_lines_for_tpl(q: dict[str, Any]) -> list[str]:
    """Формирует блок «Кому» для шапки делового письма.

    Порядок строк (по ГОСТ Р 7.0.97-2016):
    1. Должность адресата в дательном падеже
    2. Наименование организации
    3. Адрес
    4. Фамилия в дательном падеже + инициалы (напр. «Иванову И.И.»)
    """
    lines: list[str] = []
    pos = (q.get("recipient_contact_position") or "").strip()
    if pos:
        lines.append(pos)
    rn = (q.get("recipient_name") or "").strip()
    if rn:
        for part in rn.splitlines():
            p = part.strip()
            if p:
                lines.append(p)
    ra = (q.get("recipient_address") or "").strip()
    if ra:
        for part in ra.splitlines():
            p = part.strip()
            if p:
                lines.append(p)
    cn = (q.get("recipient_contact_name") or "").strip()
    if cn:
        lines.append(cn)
    cp = (q.get("recipient_contact_phone") or "").strip()
    if cp:
        lines.append(f"тел.: {cp}")
    ce = (q.get("recipient_contact_email") or "").strip()
    if ce:
        lines.append(f"e-mail: {ce}")
    return lines


def _price_header_currency(currency: str | None) -> str:
    c = (currency or "RUB").upper()
    if c in ("KZT", "KZTEN"):
        return "Цена, тенге с НДС"
    if c in ("RUB", "RUR"):
        return "Цена, руб. с НДС"
    return "Цена с НДС"


def _sum_header_currency(currency: str | None) -> str:
    c = (currency or "RUB").upper()
    if c in ("KZT", "KZTEN"):
        return "Стоимость, тенге с НДС"
    if c in ("RUB", "RUR"):
        return "Стоимость, руб. с НДС"
    return "Стоимость с НДС"


def _disc_header_currency(currency: str | None) -> str:
    c = (currency or "RUB").upper()
    if c in ("KZT", "KZTEN"):
        return "Стоимость со скидкой, тенге с НДС"
    if c in ("RUB", "RUR"):
        return "Стоимость со скидкой, руб. с НДС"
    return "Стоимость со скидкой с НДС"


# ──────────────────────────────────────────────────────────────────────────────
# Загрузка фото / логотипа / печати → InlineImage
# ──────────────────────────────────────────────────────────────────────────────

import re as _re
import http.cookiejar as _cookiejar

_GDRIVE_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    )
}


def _gdrive_file_id(url: str) -> str | None:
    """Извлекает FILE_ID из любой ссылки Google Drive."""
    m = _re.search(r"/file/d/([A-Za-z0-9_-]+)", url)
    if m:
        return m.group(1)
    m = _re.search(r"[?&]id=([A-Za-z0-9_-]+)", url)
    if m:
        return m.group(1)
    return None


def _fetch_gdrive_bytes(file_id: str) -> bytes:
    """Скачивает файл из Google Drive по file_id, обходя страницу подтверждения.

    Использует два метода:
    1. drive.usercontent.google.com — новый прямой URL (работает для публичных файлов)
    2. drive.google.com/uc с обработкой куков и confirm-токена (запасной путь)
    """
    # Метод 1: новый URL (drive.usercontent)
    usercontent_url = (
        f"https://drive.usercontent.google.com/download"
        f"?id={file_id}&export=download&authuser=0"
    )
    try:
        cookie_jar = _cookiejar.CookieJar()
        opener = urllib.request.build_opener(
            urllib.request.HTTPCookieProcessor(cookie_jar)
        )
        opener.addheaders = list(_GDRIVE_HEADERS.items())
        with opener.open(usercontent_url, timeout=25) as resp:
            ct = resp.headers.get("Content-Type", "")
            raw = resp.read()
        if raw and "text/html" not in ct:
            return raw
    except Exception:
        pass

    # Метод 2: drive.google.com/uc с confirm-токеном
    base_url = f"https://drive.google.com/uc?export=download&id={file_id}"
    cookie_jar = _cookiejar.CookieJar()
    opener = urllib.request.build_opener(
        urllib.request.HTTPCookieProcessor(cookie_jar)
    )
    opener.addheaders = list(_GDRIVE_HEADERS.items())

    with opener.open(base_url, timeout=25) as resp:
        ct = resp.headers.get("Content-Type", "")
        raw = resp.read()

    if raw and "text/html" not in ct:
        return raw

    # HTML — ищем confirm-токен в куках или теле
    confirm = None
    for cookie in cookie_jar:
        if cookie.name.startswith("download_warning"):
            confirm = cookie.value
            break
    if not confirm:
        m = _re.search(rb"confirm=([0-9A-Za-z_\-]+)", raw)
        if m:
            confirm = m.group(1).decode()
    if not confirm:
        m = _re.search(rb'["\']confirm["\']\s*:\s*["\']([^"\']+)["\']', raw)
        if m:
            confirm = m.group(1).decode()

    if confirm:
        confirm_url = f"{base_url}&confirm={confirm}"
        with opener.open(confirm_url, timeout=25) as resp2:
            raw = resp2.read()

    return raw


_UPLOADS_DIR = Path(__file__).resolve().parent.parent / "uploads"


def _fetch_image(url: str, tpl: DocxTemplate, width_cm: float) -> "InlineImage | str":
    """Скачивает изображение по URL и возвращает InlineImage.

    Поддерживает:
    - Локальные пути /uploads/... → читает файл с диска (быстро, всегда работает)
    - Google Drive ссылки → конвертирует и скачивает
    - Любые http/https URL → скачивает напрямую

    Возвращает '' при любой ошибке — тег в шаблоне просто исчезает.
    """
    if not url or not url.strip():
        return ""
    url = url.strip()
    try:
        # Локальный файл — читаем с диска, без HTTP
        if url.startswith("/uploads/"):
            local_path = _UPLOADS_DIR / url[len("/uploads/"):]
            if not local_path.is_file():
                _log.warning("Локальный файл не найден: %s", local_path)
                return ""
            raw = local_path.read_bytes()
        elif "drive.google.com" in url or "drive.usercontent.google.com" in url:
            file_id = _gdrive_file_id(url)
            if not file_id:
                _log.warning("Не удалось извлечь file_id из Google Drive URL: %s", url)
                return ""
            raw = _fetch_gdrive_bytes(file_id)
        else:
            req = urllib.request.Request(url, headers=_GDRIVE_HEADERS)
            with urllib.request.urlopen(req, timeout=20) as resp:
                raw = resp.read()

        if not raw:
            return ""
        # Ensure image is in a format python-docx understands (no WebP/AVIF/etc.)
        # Re-encode via Pillow → PNG to guarantee compatibility.
        try:
            from PIL import Image as _PILImage
            buf_out = io.BytesIO()
            _PILImage.open(io.BytesIO(raw)).convert("RGB").save(buf_out, format="JPEG")
            raw = buf_out.getvalue()
        except Exception as _conv_exc:
            _log.warning("Конвертация изображения не удалась (%s): %s", url, _conv_exc)
            return ""
        return InlineImage(tpl, io.BytesIO(raw), width=Cm(width_cm))
    except Exception as exc:
        _log.warning("Изображение не загружено (%s): %s", url, exc)
        return ""


# Псевдоним для обратной совместимости (фото товаров)
def _fetch_photo_as_inline(
    url: str, tpl: DocxTemplate, width_cm: float = 14.0
) -> "InlineImage | str":
    return _fetch_image(url, tpl, width_cm)


# ──────────────────────────────────────────────────────────────────────────────
# Сборка контекста Jinja2
# ──────────────────────────────────────────────────────────────────────────────

def build_docxtpl_context(ctx: dict[str, Any]) -> dict[str, Any]:
    """Возвращает словарь переменных для рендера шаблона.

    Поле item['photo'] выставляется в '' — build_quote_docxtpl() заменит его
    на InlineImage после создания объекта DocxTemplate.
    """
    q = ctx["quote"]
    sender = ctx.get("sender") or {}
    items_raw = ctx.get("items") or []
    show_disc = q.get("show_discount_column", True)

    # Регулярки для авто-разбивки intro на пули и текст описания
    _BULLET_RE = _re.compile(r'^[•\-–—*]|^\d+[).]\s')
    _LABEL_RE  = _re.compile(
        r'^(Главн[аые]+\s+особенност[ьи]|Описание\s+товара|'
        r'Применение|Назначение|Комплект\s+поставки|'
        r'Основные\s+характеристики|Преимущества)\s*$', _re.I
    )

    items_out: list[dict[str, Any]] = []
    for idx, it in enumerate(items_raw, start=1):
        row = _normalize_item_export(it)
        ft = (row.get("features_text") or "").strip()
        features_lines = [x.strip() for x in ft.splitlines() if x.strip()] if ft else []
        specs_in = row.get("specs") or []
        spec_pairs: list[dict[str, str]] = []
        for s in specs_in:
            if not isinstance(s, dict):
                continue
            k = str(s.get("param") or s.get("name") or s.get("key") or "").strip()
            v = str(s.get("value") or s.get("val") or "").strip()
            if k or v:
                spec_pairs.append({"param": k, "value": v})
        title = (row.get("title") or "").strip()
        model = (row.get("model") or "").strip()
        cap_title = f"{title} Модель: {model}".strip() if model else title

        # Предварительно форматированные блоки — Listing сохраняет \n как
        # <w:br/> в XML. Используются в {{ item.features_block }} и
        # {{ item.specs_block }} вместо вложенных {% for %} (которые ломаются
        # из-за разрыва Jinja2-тегов XML-элементом <w:br/>).
        features_block: "Listing | str" = (
            Listing("\n".join(features_lines))
            if features_lines else ""
        )
        specs_block: "Listing | str" = (
            Listing("\n".join(f"{sp['param']}: {sp['value']}" for sp in spec_pairs))
            if spec_pairs else ""
        )

        intro_plain = (row.get("intro") or "").strip()
        kit_plain = (row.get("kit_text") or "").strip()

        # Если features_text пустой, но в intro хранятся пули вперемешку
        # с текстом описания — автоматически разделяем их.
        # Строки, начинающиеся с «•», «–», «-» или цифры+точки → особенности.
        # Всё остальное (кроме коротких заголовков-меток) → описание.
        if not features_lines and intro_plain:
            _desc_parts: list[str] = []
            _feat_parts: list[str] = []
            for _ln in intro_plain.splitlines():
                _s = _ln.strip()
                if not _s or _LABEL_RE.match(_s):
                    continue                       # всегда удаляем заголовки-метки
                if _BULLET_RE.match(_s):
                    _feat_parts.append(_s)         # пуля → особенности
                else:
                    _desc_parts.append(_s)         # текст → описание
            if _feat_parts:
                features_lines = _feat_parts
                features_block = Listing("\n".join(features_lines))
            intro_plain = "\n".join(_desc_parts).strip()  # всегда обновляем

        _show_features = bool(row.get("show_features", True))
        _show_intro    = bool(row.get("show_intro",    True))
        # Если флаг отключён — пустим пустые списки, чтобы {%p for %} не рендерил строки
        if not _show_features:
            features_lines = []
            features_block = ""

        items_out.append(
            {
                "title": title,
                "model": model,
                "show_intro": _show_intro,
                "show_features": _show_features,
                "show_kit": row.get("show_kit", True),
                "show_specs": row.get("show_specs", True),
                "show_photos": row.get("show_photos", True),
                "intro": intro_plain,
                "intro_listing": Listing(intro_plain) if intro_plain else "",
                "kit_text": kit_plain,
                "kit_listing": Listing(kit_plain) if kit_plain else "",
                "features_lines": features_lines,
                "features_block": features_block,
                "specs": spec_pairs,
                "specs_block": specs_block,
                "figure_caption": f"Рис.{idx} {cap_title}",
                "photo_urls": list(row.get("photo_urls") or []),
                "photo": "",
            }
        )

    spec_rows: list[dict[str, Any]] = []
    sum_full = 0.0
    sum_disc = 0.0
    for ri, it in enumerate(items_raw, start=1):
        name = (it.get("title") or "").strip()
        if it.get("model"):
            name = f"{name} Модель: {it['model']}".strip()
        pw = float(it.get("price_with_vat") or 0)
        qty = float(it.get("quantity") or 1)
        lt = float(it.get("line_total_with_vat") or 0)
        ld = float(it.get("line_total_discounted") or 0)
        sum_full += lt
        sum_disc += ld
        qs = str(int(qty)) if float(qty).is_integer() else str(qty)
        spec_rows.append(
            {
                "no": str(ri),
                "name": name,
                "price": money_ru(pw),
                "qty": qs,
                "sum_full": money_ru(lt),
                "sum_disc": money_ru(ld),
            }
        )

    sender_lines = _sender_lines_for_tpl(sender, q.get("currency"))
    recipient_lines = _recipient_lines_for_tpl(q)
    intro_plain = (q.get("intro_text") or "").strip()

    # ── Отдельные поля отправителя для kp_rutesttpl (пер-строчное форматирование) ──
    _currency = q.get("currency")
    _tl = _tax_label_for_currency(_currency)
    _lf   = _expand_org_form((sender.get("legal_form") or "").strip())
    _ln   = (sender.get("legal_name") or sender.get("name") or "").strip()
    _addr = (sender.get("legal_address") or "").strip()
    _inn  = (sender.get("tax_number") or "").strip()
    _kpp  = (sender.get("kpp") or "").strip()
    _ogrn = (sender.get("ogrn") or "").strip()
    _em   = (sender.get("email") or "").strip()
    _ph   = (sender.get("phone") or "").strip()
    _web  = (sender.get("website") or "").strip()
    _inn_kpp_parts: list[str] = []
    if _inn:
        _inn_kpp_parts.append(f"{_tl} {_inn}")
    if _kpp and _tl == "ИНН":
        _inn_kpp_parts.append(f"КПП {_kpp}")
    _inn_kpp_str = (
        ", ".join(_inn_kpp_parts) + ("," if _kpp and _tl == "ИНН" else "")
        if _inn_kpp_parts else ""
    )

    return {
        # Исходные списки — для пользовательских шаблонов с {% for line in sender_lines %}
        "sender_lines": sender_lines,
        "recipient_lines": recipient_lines,
        # Готовый многострочный текст через Listing — для встроенного шаблона
        "sender_text": Listing("\n".join(sender_lines)) if sender_lines else "",
        "recipient_text": Listing("\n".join(recipient_lines)) if recipient_lines else "",
        # ── Отдельные поля отправителя (kp_rutesttpl.docx — пер-строчное форм.) ──
        "sender_org_form":   _lf,
        "sender_legal_name": f"«{_ln}»" if _ln else "",
        "sender_address":    _addr,
        "sender_inn_kpp":    _inn_kpp_str,
        "sender_ogrn":       f"ОГРН {_ogrn}" if _ogrn else "",
        "sender_email":      f"e-mail: {_em}" if _em else "",
        "sender_phone":      f"Тел.: {_ph}" if _ph else "",
        "sender_website":    f"Сайт: {_web}" if _web else "",
        # ── Отдельные поля получателя (kp_rutesttpl.docx) ──────────────────────
        "recipient_role":    (q.get("recipient_contact_position") or "").strip(),
        "recipient_org":     (q.get("recipient_name") or "").strip(),
        "recipient_person":  (q.get("recipient_contact_name") or "").strip(),
        "quote_date_ru": format_quote_date_ru(q.get("quote_date")),
        "quote_number": q.get("number") or "",
        # intro — строка для {% if intro %}; intro_listing — многострочный текст с переносами
        "intro": intro_plain,
        "intro_listing": Listing(intro_plain) if intro_plain else "",
        "greeting_name": (q.get("greeting_name") or "").strip(),
        "items_short_list": _items_short_listing(items_raw),
        "items": items_out,
        "spec_rows": spec_rows,
        "show_discount_column": show_disc,
        "sum_total_full": money_ru(sum_full),
        "sum_total_disc": money_ru(sum_disc),
        "price_col_title": _price_header_currency(q.get("currency")),
        "sum_col_title": _sum_header_currency(q.get("currency")),
        "disc_col_title": _disc_header_currency(q.get("currency")),
        "terms_price_validity": (q.get("terms_price_validity") or "").strip(),
        "terms_delivery": (q.get("terms_delivery") or "").strip(),
        "terms_lead_time": (q.get("terms_lead_time") or "").strip(),
        "terms_payment": (q.get("terms_payment") or "").strip(),
        "terms_production_country": (q.get("terms_production_country") or "").strip(),
        "terms_warranty": (q.get("terms_warranty") or "").strip(),
        "terms_currency_note": (
            (q.get("terms_currency_note") or "").strip()
            or _currency_note(q.get("currency"))
        ),
        "terms_address_note": (q.get("terms_address_note") or "").strip(),
        "vat_rate": q.get("vat_rate") or 20,
        "vat_amount": money_ru(max(sum_full - sum_disc, sum_full) * (q.get("vat_rate") or 20) / (100 + (q.get("vat_rate") or 20))),
        "sender": sender,
        "signer_name": (sender.get("signer_name") or "").strip(),
        "signer_position": (sender.get("signer_position") or "").strip(),
        "signer_phone": (sender.get("phone") or "").strip(),
        # Изображения — заглушки ""; реальные InlineImage вставляются в build_quote_docxtpl()
        "logo": "",
        "signature": "",
        "stamp": "",
    }


# ──────────────────────────────────────────────────────────────────────────────
# Встроенный минимальный шаблон
# ──────────────────────────────────────────────────────────────────────────────

def _write_builtin_docxtpl_template(path: Path) -> None:
    """Генерирует минимальный .docx-шаблон с корректными тегами docxtpl.

    Правила встроенного шаблона (отличаются от пользовательского):
    - Никаких \\n внутри run-текста, содержащего {%...%}: python-docx превращает
      \\n в <w:br/>, что разбивает Jinja2-теги и вызывает TemplateSyntaxError.
    - Вместо {% for line in sender_lines %} используем {{ sender_text }} (Listing).
    - Вместо вложенных {% for %} используем {{ item.features_block }} / specs_block.
    - Спецификация: 4-строчная таблица ({%tr for %} и {%tr endfor %} — отдельные
      строки, потому что {%tr %} заменяет ВСЮ строку таблицы целиком).

    Для production используйте свой Word-шаблон с любым оформлением —
    передайте путь через QUOTE_DOCX_TEMPLATE_PATH.
    """
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)

    # ── Шапка: отправитель | получатель ──────────────────────────────
    # Используем {{ sender_text }} (Listing) вместо цикла с \n,
    # чтобы не создавать <w:br/> внутри Jinja2-тегов.
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.cell(0, 0).paragraphs[0].add_run("{{ sender_text }}")
    tbl.cell(0, 1).paragraphs[0].add_run("{{ recipient_text }}")

    # ── Дата и заголовок ──────────────────────────────────────────────
    pd = doc.add_paragraph()
    pd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pd.add_run("{{ quote_date_ru }}")

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = pt.add_run("Коммерческое предложение")
    run_title.bold = True
    run_title.font.size = Pt(14)

    pnum = doc.add_paragraph()
    pnum.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pnum.add_run("{% if quote_number %}№ {{ quote_number }}{% endif %}")

    pgreet = doc.add_paragraph()
    # Условие БЕЗ \n внутри тегов — вся строка на одной строке шаблона
    pgreet.add_run(
        "{% if intro %}{{ intro_listing }}{% elif greeting_name %}"
        "Уважаемый {{ greeting_name }}!{% else %}Уважаемый заказчик!{% endif %}"
    )

    doc.add_paragraph("")
    ptech = doc.add_paragraph()
    ptech.add_run("ТЕХНИЧЕСКОЕ ОПИСАНИЕ ПРИБОРОВ").bold = True
    doc.add_paragraph("")

    # ── Блоки позиций ─────────────────────────────────────────────────
    # ПРАВИЛО: NO \n внутри тела {% for %} в одном run.
    # Многострочные поля (features_block, specs_block) — Listing-объекты,
    # которые docxtpl сам преобразует в <w:br/> при рендере.
    # item.photo — InlineImage или "" (тег исчезнет, если фото недоступно).
    pb = doc.add_paragraph()
    pb.add_run(
        "{% for item in items %}"
        "{{ item.title }}"
        "{% if item.model %} / Модель: {{ item.model }}{% endif %}"
        "{% if item.intro %} | {{ item.intro }}{% endif %}"
        "{% if item.kit_text %} | Комплект: {{ item.kit_text }}{% endif %}"
        "{% if item.features_block %} | Особенности: {{ item.features_block }}{% endif %}"
        "{% if item.specs_block %} | ТХ: {{ item.specs_block }}{% endif %}"
        " | {{ item.figure_caption }}"
        " | {{ item.photo }}"
        " ~~~ "
        "{% endfor %}"
    )
    doc.add_paragraph(
        "(Для красивого оформления позиций загрузите свой шаблон Word "
        "через QUOTE_DOCX_TEMPLATE_PATH)"
    )

    # ── Спецификация ──────────────────────────────────────────────────
    # {%tr for r in spec_rows %} заменяет ВСЮ строку таблицы целиком.
    # Поэтому нужны 4 отдельные строки:
    #   [0] Заголовки
    #   [1] Строка-тег {%tr for r in spec_rows %} — весь ряд становится {% for %}
    #   [2] Строка данных {{ r.no }}, {{ r.name }}, ... — повторяется N раз
    #   [3] Строка-тег {%tr endfor %}              — весь ряд становится {% endfor %}
    doc.add_paragraph("")
    hs = doc.add_paragraph()
    hs.add_run("СПЕЦИФИКАЦИЯ НА ПОСТАВКУ ТОВАРОВ").bold = True
    doc.add_paragraph("")

    spec_tbl = doc.add_table(rows=4, cols=6)
    spec_tbl.style = "Table Grid"

    # Строка 0: заголовки
    for j, txt in enumerate([
        "№", "Наименование",
        "{{ price_col_title }}", "Кол-во",
        "{{ sum_col_title }}", "{{ disc_col_title }}",
    ]):
        spec_tbl.cell(0, j).paragraphs[0].add_run(txt).bold = True

    # Строка 1: тег цикла — вся строка заменится на {% for r in spec_rows %}
    spec_tbl.cell(1, 0).paragraphs[0].add_run("{%tr for r in spec_rows %}")

    # Строка 2: содержимое — будет повторяться для каждой строки спецификации
    for j, txt in enumerate([
        "{{ r.no }}", "{{ r.name }}", "{{ r.price }}",
        "{{ r.qty }}", "{{ r.sum_full }}", "{{ r.sum_disc }}",
    ]):
        spec_tbl.cell(2, j).paragraphs[0].add_run(txt)

    # Строка 3: тег endfor — вся строка заменится на {% endfor %}
    spec_tbl.cell(3, 0).paragraphs[0].add_run("{%tr endfor %}")

    # Итого после таблицы
    doc.add_paragraph("")
    tot = doc.add_paragraph()
    tot.add_run(
        "ИТОГО: {{ sum_total_full }}"
        "{% if show_discount_column %} | Итого со скидкой: {{ sum_total_disc }}{% endif %}"
    )

    # ── Условия поставки ──────────────────────────────────────────────
    doc.add_paragraph("")
    for tmpl in [
        "{% if terms_price_validity %}{{ terms_price_validity }}{% endif %}",
        "{% if terms_production_country %}Страна: {{ terms_production_country }}{% endif %}",
        "{% if terms_warranty %}Гарантия: {{ terms_warranty }}{% endif %}",
        "{% if terms_delivery %}Условия поставки: {{ terms_delivery }}{% endif %}",
        "{% if terms_lead_time %}Срок поставки: {{ terms_lead_time }}{% endif %}",
        "{% if terms_payment %}Условия оплаты: {{ terms_payment }}{% endif %}",
        "{{ terms_currency_note }}",
        "{% if terms_address_note %}{{ terms_address_note }}{% endif %}",
    ]:
        doc.add_paragraph(tmpl)

    # ── Подпись ───────────────────────────────────────────────────────
    doc.add_paragraph("")
    doc.add_paragraph("С уважением,")
    doc.add_paragraph("{{ signer_name }}")
    doc.add_paragraph("{{ signer_position }}")
    doc.add_paragraph("{{ signer_phone }}")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


# ──────────────────────────────────────────────────────────────────────────────
# Публичный API
# ──────────────────────────────────────────────────────────────────────────────

def resolve_docxtpl_template_path() -> Path:
    env = os.environ.get("QUOTE_DOCX_TEMPLATE_PATH", "").strip()
    root = Path(__file__).resolve().parent.parent
    if env:
        p = Path(env).expanduser()
        if not p.is_absolute():
            p = root / p
        return p.resolve()
    default = root / "templates" / "kp_docxtpl.docx"
    if not default.exists():
        _write_builtin_docxtpl_template(default)
    return default


def build_quote_docxtpl(ctx: dict[str, Any]) -> bytes:
    """Рендерит КП через docxtpl и возвращает байты .docx.

    Порядок работы:
    1. Открывает шаблон (DocxTemplate).
    2. Собирает Jinja2-контекст через build_docxtpl_context().
    3. Для каждой позиции с show_photos=True и непустыми photo_urls —
       скачивает первое фото и подставляет InlineImage в item['photo'].
       При недоступности URL поле остаётся ''.
    4. Рендерит и возвращает байты.
    """
    tpl_path = resolve_docxtpl_template_path()
    if not tpl_path.is_file():
        raise FileNotFoundError(f"Шаблон КП не найден: {tpl_path}")

    tpl = DocxTemplate(str(tpl_path))
    context = build_docxtpl_context(ctx)

    # Логотип, подпись, печать из профиля отправителя
    sender = ctx.get("sender") or {}
    context["logo"]      = _fetch_image(sender.get("logo_url", ""),      tpl, width_cm=3.0)
    context["signature"] = _fetch_image(sender.get("signature_url", ""), tpl, width_cm=3.5)
    context["stamp"]     = _fetch_image(sender.get("stamp_url", ""),     tpl, width_cm=3.5)

    # Фото товаров — InlineImage привязан к tpl
    raw_items = ctx.get("items") or []
    for i, item_ctx in enumerate(context["items"]):
        photo: Any = ""
        if i < len(raw_items):
            raw = raw_items[i]
            urls = list(raw.get("photo_urls") or [])
            if raw.get("show_photos") and urls:
                photo = _fetch_image(urls[0], tpl, width_cm=14.0)
        item_ctx["photo"] = photo

    tpl.render(context)

    # Сохраняем рендер в байты
    buf = io.BytesIO()
    tpl.save(buf)

    # ── Постобработка документа ──────────────────────────────────────────────
    try:
        import random as _random
        from docx import Document as _Document
        from lxml import etree as _et
        from copy import deepcopy as _deepcopy
        _W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        _W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
        _P   = f"{{{_W}}}p"
        _R   = f"{{{_W}}}r"
        _RPR = f"{{{_W}}}rPr"
        _BR  = f"{{{_W}}}br"
        _PPR = f"{{{_W}}}pPr"
        _JC  = f"{{{_W}}}jc"
        _IND = f"{{{_W}}}ind"
        _T   = f"{{{_W}}}t"
        _VAL = f"{{{_W}}}val"
        _FL  = f"{{{_W}}}firstLine"
        _TC  = f"{{{_W}}}tc"
        _XML_SP = "{http://www.w3.org/XML/1998/namespace}space"
        _PARA_ID  = f"{{{_W14}}}paraId"
        _TEXT_ID  = f"{{{_W14}}}textId"

        def _in_table_cell(el):
            p = el.getparent()
            while p is not None:
                if p.tag == _TC:
                    return True
                p = p.getparent()
            return False

        buf.seek(0)
        _rdoc = _Document(buf)

        # ── Проход 0: устраняем дублирующиеся ID ────────────────────────────
        # docxtpl клонирует параграфы/объекты шаблона при итерации, не обновляя
        # уникальные атрибуты. Word требует уникальности и показывает диалог
        # «восстановить содержимое» при дублях.

        # 0a) w14:paraId / w14:textId на параграфах
        _seen_para_ids: set = set()
        _seen_text_ids: set = set()
        for _para in _rdoc.element.iter(_P):
            _pid = _para.get(_PARA_ID)
            if _pid is not None:
                if _pid in _seen_para_ids:
                    _para.set(_PARA_ID, f"{_random.randint(1, 0x7FFFFFFE):08X}")
                else:
                    _seen_para_ids.add(_pid)
            _tid = _para.get(_TEXT_ID)
            if _tid is not None:
                if _tid in _seen_text_ids:
                    _para.set(_TEXT_ID, f"{_random.randint(1, 0x7FFFFFFE):08X}")
                else:
                    _seen_text_ids.add(_tid)

        # 0b) id на <wp:cNvPr> и <pic:cNvPr> (InlineImage / Drawing objects)
        # Все картинки должны иметь уникальный числовой id >= 1.
        _WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        _PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
        _A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
        _seen_drawing_ids: set = set()
        _next_drawing_id = [1]

        def _new_drawing_id():
            while _next_drawing_id[0] in _seen_drawing_ids:
                _next_drawing_id[0] += 1
            new_id = _next_drawing_id[0]
            _seen_drawing_ids.add(new_id)
            _next_drawing_id[0] += 1
            return str(new_id)

        for _cnv in _rdoc.element.iter(
            f"{{{_WP}}}cNvPr",
            f"{{{_PIC}}}cNvPr",
        ):
            _did = _cnv.get("id")
            if _did is None:
                continue
            try:
                _did_int = int(_did)
            except ValueError:
                continue
            if _did_int < 1 or _did_int in _seen_drawing_ids:
                _cnv.set("id", _new_drawing_id())
            else:
                _seen_drawing_ids.add(_did_int)

        # ── Проход 0c: пустые ячейки таблицы — добавляем пустой w:p ────────
        # OOXML требует, чтобы каждый <w:tc> содержал хотя бы один <w:p>.
        # Если {%p for %} вернул 0 строк (пустой список), ячейка остаётся
        # без параграфов — Word выдаёт «содержимое, которое не удалось прочитать».
        for _tc in _rdoc.element.iter(_TC):
            if not any(c.tag == _P for c in _tc):
                _empty_p = _et.SubElement(_tc, _P)

        # ── Проход 1: исправляем justify у br-абзацев ───────────────────────
        for _para in _rdoc.element.iter(_P):
            _pPr = _para.find(_PPR)
            if any(el.tag == _BR for el in _para.iter()):
                if _pPr is not None:
                    _jc = _pPr.find(_JC)
                    if _jc is not None and _jc.get(_VAL) in ("both", "distribute"):
                        _jc.set(_VAL, "left")

        # ── Проход 2: разбиваем Listing br-абзацы вне таблиц на отдельные w:p ──
        # Каждый «логический» кусок текста (между <w:br/>) становится
        # отдельным абзацем, копируя w:pPr (incl. firstLine) из оригинала.
        # Это позволяет w:firstLine работать для КАЖДОГО пункта списка.
        for _para in list(_rdoc.element.iter(_P)):
            if _in_table_cell(_para):
                continue
            _pPr = _para.find(_PPR)
            # Собираем потомков первого уровня (не рекурсивно) для проверки
            if not any(el.tag == _BR for el in _para.iter()):
                continue

            _parent = _para.getparent()
            if _parent is None:
                continue
            _idx = list(_parent).index(_para)

            # Разбираем содержимое: обходим дочерние элементы <w:p>
            # и строим список «строк» (segments), разделённых <w:br/>.
            # Структура docxtpl Listing: один <w:r> с чередованием <w:t>/<w:br/>.
            _segments: list[tuple] = []  # list of (rpr_el, text_str)
            _cur_rpr = None
            _cur_parts: list[str] = []

            def _flush():
                txt = "".join(_cur_parts)
                if txt.strip():
                    _segments.append((_cur_rpr, txt))
                _cur_parts.clear()

            for _child in list(_para):
                if _child.tag == _PPR:
                    continue
                if _child.tag == _R:
                    for _gc in list(_child):
                        if _gc.tag == _RPR:
                            _cur_rpr = _gc
                        elif _gc.tag == _T:
                            _cur_parts.append(_gc.text or "")
                        elif _gc.tag == _BR:
                            _flush()
                elif _child.tag == _T:
                    _cur_parts.append(_child.text or "")
                elif _child.tag == _BR:
                    _flush()
            _flush()  # последний кусок

            if len(_segments) <= 1:
                continue  # нечего разбивать

            # Создаём новые <w:p> — по одному на каждый непустой сегмент
            _new_paras = []
            for _rpr, _txt in _segments:
                _np = _et.Element(_P)
                if _pPr is not None:
                    _np.append(_deepcopy(_pPr))
                _nr = _et.SubElement(_np, _R)
                if _rpr is not None:
                    _nr.append(_deepcopy(_rpr))
                _nt = _et.SubElement(_nr, _T)
                if _txt != _txt.strip():
                    _nt.set(_XML_SP, "preserve")
                _nt.text = _txt.strip()
                _new_paras.append(_np)

            # Вставляем новые абзацы и удаляем оригинал
            for _i, _np in enumerate(_new_paras):
                _parent.insert(_idx + _i, _np)
            _parent.remove(_para)

        # ── Проход 3: красная строка всем обычным параграфам ────────────────
        # Применяется только к параграфам:
        #   - вне таблиц
        #   - без явного центрирования (jc=center)
        #   - без явного right-выравнивания
        #   - которые содержат текст (не пустые)
        # Размер красной строки: 720 двадцатых пункта (~1.27 см).
        _INDENT_TWENTIETHS = "720"  # 1.27 см
        for _para in _rdoc.element.iter(_P):
            if _in_table_cell(_para):
                continue
            # Текст в параграфе?
            _has_text = False
            for _t in _para.iter(_T):
                if (_t.text or "").strip():
                    _has_text = True
                    break
            if not _has_text:
                continue
            # Получаем/создаём pPr
            _pPr = _para.find(_PPR)
            if _pPr is None:
                _pPr = _et.SubElement(_para, _PPR)
                _para.insert(0, _pPr)
                _para.remove(_pPr)
                _para.insert(0, _pPr)
            # Пропускаем центрированные/правые
            _jc = _pPr.find(_JC)
            if _jc is not None and _jc.get(_VAL) in ("center", "right", "end"):
                continue
            # Уже задан firstLine или hanging? — не перезаписываем
            _ind = _pPr.find(_IND)
            if _ind is not None and (_ind.get(_FL) or _ind.get(f"{{{_W}}}hanging")):
                continue
            if _ind is None:
                _ind = _et.SubElement(_pPr, _IND)
            _ind.set(_FL, _INDENT_TWENTIETHS)

        buf2 = io.BytesIO()
        _rdoc.save(buf2)
        return buf2.getvalue()
    except Exception as _pp_exc:
        _log.warning("Постобработка не удалась: %s", _pp_exc)
        buf.seek(0)
        return buf.read()
