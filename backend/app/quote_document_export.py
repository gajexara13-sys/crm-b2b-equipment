"""Экспорт коммерческого предложения в DOCX и PDF (кириллица).

Макет ориентирован на типовые КП: двухколоночная шапка (отправитель / получатель),
центрированный заголовок, вводный текст, блоки по позициям (описание, особенности,
технические характеристики строками), спецификация, условия и подпись.
"""

from __future__ import annotations

import io
import os
import re
import urllib.error
import urllib.request
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

# Подсказки в шапке, если не заполнены реквизиты (вместо символа «—»)
QUOTE_HEADER_HINT_SENDER = (
    "Выберите в карточке КП профиль отправителя и заполните реквизиты организации "
    "(юр. лицо, адрес, БИН/ИНН, логотип по URL) — тогда шапка совпадёт с вашим Word-шаблоном."
)
QUOTE_HEADER_HINT_RECIPIENT = (
    "Укажите получателя: поле «Получатель (текст)» или компания из справочника, адрес и контакт."
)

_MONTHS_RU = (
    "января",
    "февраля",
    "марта",
    "апреля",
    "мая",
    "июня",
    "июля",
    "августа",
    "сентября",
    "октября",
    "ноября",
    "декабря",
)


def _parse_date_any(val: Any) -> date | None:
    if val is None:
        return None
    if isinstance(val, date):
        return val
    s = str(val).strip()
    if not s:
        return None
    try:
        return date.fromisoformat(s[:10])
    except ValueError:
        return None


def format_quote_date_ru(val: Any) -> str:
    d = _parse_date_any(val)
    if not d:
        return ""
    return f"{d.day} {_MONTHS_RU[d.month - 1]} {d.year} г."


def money_ru(x: float | None) -> str:
    v = float(x or 0)
    s = f"{v:,.2f}"
    intp, frac = s.split(".")
    intp = intp.replace(",", " ")
    return f"{intp},{frac}"


def _currency_note(code: str | None) -> str:
    c = (code or "RUB").upper()
    if c in ("RUB", "RUR"):
        return "Цены указаны в рублях."
    if c in ("KZT", "KZTEN"):
        return "Цены указаны в тенге."
    return f"Цены указаны в {c}."


def _esc_xml(s: str) -> str:
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


_UPLOADS_DIR = Path(__file__).resolve().parent.parent / "uploads"


def _fetch_url_bytes(url: str | None, *, timeout: float = 18.0, max_bytes: int = 12_000_000) -> bytes | None:
    if not url:
        return None
    url = str(url).strip()
    # Локальный файл из /uploads/ — читаем с диска без HTTP
    if url.startswith("/uploads/"):
        local = _UPLOADS_DIR / url[len("/uploads/"):]
        if local.is_file():
            data = local.read_bytes()
            return data if len(data) <= max_bytes else data[:max_bytes]
        return None
    if not url.startswith(("http://", "https://")):
        return None
    req = urllib.request.Request(url, headers={"User-Agent": "CRM-KP-Export/1.0"})
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            if getattr(resp, "status", 200) != 200:
                return None
            data = resp.read()
            return data if len(data) <= max_bytes else data[:max_bytes]
    except (urllib.error.URLError, TimeoutError, OSError, ValueError):
        return None


def _sender_styled_blocks(sender: dict[str, Any]) -> list[tuple[str, bool]]:
    """Текст левой колонки шапки: (строка, жирная ли короткая строка «название»)."""
    blocks: list[tuple[str, bool]] = []
    lf = (sender.get("legal_form") or "").strip()
    ln = (sender.get("legal_name") or sender.get("name") or "").strip()
    if lf:
        blocks.append((lf, False))
    if ln:
        blocks.append((f"«{ln}»", True))
    addr = (sender.get("legal_address") or "").strip()
    if addr:
        blocks.append((addr, False))
    inn = (sender.get("tax_number") or "").strip()
    kpp = (sender.get("kpp") or "").strip()
    ogrn = (sender.get("ogrn") or "").strip()
    tax_parts = []
    if inn:
        tax_parts.append(f"ИНН {inn}")
    if kpp:
        tax_parts.append(f"КПП {kpp}")
    if tax_parts:
        blocks.append((", ".join(tax_parts) + ("," if kpp else ""), False))
    if ogrn:
        blocks.append((f"ОГРН {ogrn}", False))
    em = (sender.get("email") or "").strip()
    if em:
        blocks.append((f"e-mail: {em}", False))
    ph = (sender.get("phone") or "").strip()
    if ph:
        blocks.append((f"Тел.: {ph}", False))
    web = (sender.get("website") or "").strip()
    if web:
        blocks.append((f"Сайт: {web}", False))
    return blocks


def _recipient_styled_blocks(q: dict[str, Any]) -> list[tuple[str, bool]]:
    lines = _recipient_lines(q)
    blocks: list[tuple[str, bool]] = []
    for i, L in enumerate(lines):
        blocks.append((L, i == 0))
    return blocks


def _docx_fill_cell_styled(cell, blocks: list[tuple[str, bool]], *, body_pt: int = 12, logo_bytes: bytes | None = None):
    """Заполняет ячейку таблицы абзацами; опционально сверху — логотип."""
    cell.text = ""
    first_para = True
    if logo_bytes:
        p = cell.paragraphs[0]
        try:
            r_img = p.add_run()
            r_img.add_picture(io.BytesIO(logo_bytes), width=Cm(3.4))
        except Exception:
            pass
        p.paragraph_format.space_after = Pt(4)
        first_para = False

    for text, bold in blocks:
        if not text:
            continue
        if first_para:
            p = cell.paragraphs[0]
            first_para = False
        else:
            p = cell.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(body_pt)
        p.paragraph_format.space_after = Pt(2)


def _docx_apply_page_setup(doc: Document):
    sec = doc.sections[0]
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)


def _docx_style_normal(doc: Document):
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)


def _docx_font_table_cells(tbl):
    """Шрифт таблицы спецификации; жирный только шапка — итог остаётся как задан выше."""
    for ri, row in enumerate(tbl.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10 if ri == 0 else 11)
                    if ri == 0:
                        r.bold = True


def _docx_add_picture_maybe(doc: Document, url: str | None, *, max_width_cm: float = 14.0):
    raw = _fetch_url_bytes(url)
    if not raw:
        return
    try:
        doc.add_picture(io.BytesIO(raw), width=Cm(max_width_cm))
    except Exception:
        return


def _flowable_image_fit(raw: bytes, max_w: float, max_h: float) -> RLImage | None:
    if not raw:
        return None
    try:
        bio = io.BytesIO(raw)
        ir = ImageReader(bio)
        iw, ih = ir.getSize()
        if iw <= 0 or ih <= 0:
            return None
        scale = min(max_w / iw, max_h / ih)
        return RLImage(io.BytesIO(raw), width=iw * scale, height=ih * scale)
    except Exception:
        return None


def _sender_lines(sender: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    lf = (sender.get("legal_form") or "").strip()
    ln = (sender.get("legal_name") or sender.get("name") or "").strip()
    if lf:
        lines.append(lf)
    if ln:
        lines.append(f"«{ln}»")
    addr = (sender.get("legal_address") or "").strip()
    if addr:
        lines.append(addr)
    inn = (sender.get("tax_number") or "").strip()
    kpp = (sender.get("kpp") or "").strip()
    ogrn = (sender.get("ogrn") or "").strip()
    tax_parts = []
    if inn:
        tax_parts.append(f"ИНН {inn}")
    if kpp:
        tax_parts.append(f"КПП {kpp}")
    if tax_parts:
        lines.append(", ".join(tax_parts) + ("," if kpp else ""))
    if ogrn:
        lines.append(f"ОГРН {ogrn}")
    em = (sender.get("email") or "").strip()
    if em:
        lines.append(f"e-mail: {em}")
    ph = (sender.get("phone") or "").strip()
    if ph:
        lines.append(f"Тел.: {ph}")
    web = (sender.get("website") or "").strip()
    if web:
        lines.append(f"Сайт: {web}")
    return lines


def _recipient_lines(q: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    rn = (q.get("recipient_name") or "").strip()
    if rn:
        for part in rn.splitlines():
            part = part.strip()
            if part:
                lines.append(part)
    ra = (q.get("recipient_address") or "").strip()
    if ra:
        for part in ra.splitlines():
            part = part.strip()
            if part:
                lines.append(part)
    cn = (q.get("recipient_contact_name") or "").strip()
    pos = (q.get("recipient_contact_position") or "").strip()
    if cn:
        lines.append(f"{cn}" + (f", {pos}" if pos else ""))
    cp = (q.get("recipient_contact_phone") or "").strip()
    if cp:
        lines.append(f"тел.: {cp}")
    ce = (q.get("recipient_contact_email") or "").strip()
    if ce:
        lines.append(f"e-mail: {ce}")
    return lines


def _spec_rows(specs: list[dict[str, Any]] | None) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []
    if not specs:
        return rows
    for row in specs:
        if not isinstance(row, dict):
            continue
        k = row.get("param") or row.get("name") or row.get("key") or ""
        v = row.get("value") or row.get("val") or ""
        k, v = str(k).strip(), str(v).strip()
        if k or v:
            rows.append((k, v))
    return rows


def _item_caption(it: dict[str, Any]) -> str:
    title = (it.get("title") or "").strip()
    model = (it.get("model") or "").strip()
    if model:
        return f"{title} ({model})".strip() if title else model
    return title or "Позиция"


def _features_numbered_docx(doc: Document, features_text: str | None):
    if not features_text or not str(features_text).strip():
        return
    lines = [ln.strip() for ln in str(features_text).splitlines() if ln.strip()]
    for i, ln in enumerate(lines, start=1):
        p = doc.add_paragraph(f"{i}. {ln}")
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)


def _para_lines_docx(doc: Document, text: str | None, *, size: int = 12, justify: bool = False):
    if not text or not str(text).strip():
        return
    for line in str(text).splitlines():
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph(line)
        if justify:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(size)


def _spec_lines_docx(doc: Document, specs: list[tuple[str, str]]):
    if not specs:
        return
    p = doc.add_paragraph()
    r = p.add_run("ТЕХНИЧЕСКИЕ ХАРАКТЕРИСТИКИ")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    for k, v in specs:
        if k and v:
            line = f"{k}: {v}"
        elif k:
            line = k
        else:
            line = v
        _para_lines_docx(doc, line)


def _footer_terms_docx(doc: Document, q: dict[str, Any]):
    pairs: list[tuple[str, str | None]] = [
        ("", q.get("terms_price_validity")),
        ("Условия поставки", q.get("terms_delivery")),
        ("Срок поставки", q.get("terms_lead_time")),
        ("Условия оплаты", q.get("terms_payment")),
        ("Страна производства", q.get("terms_production_country")),
        ("Гарантия", q.get("terms_warranty")),
    ]
    for label, val in pairs:
        if not val:
            continue
        if label:
            p = doc.add_paragraph(f"{label}: {val}")
        else:
            p = doc.add_paragraph(str(val))
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
    cur = (q.get("terms_currency_note") or "").strip() or _currency_note(q.get("currency"))
    if cur:
        p = doc.add_paragraph(cur)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
    addr_note = (q.get("terms_address_note") or "").strip()
    if addr_note:
        p = doc.add_paragraph(addr_note)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)


def build_quote_docx(ctx: dict[str, Any]) -> bytes:
    q = ctx["quote"]
    sender = ctx.get("sender") or {}
    items = ctx.get("items") or []

    doc = Document()
    _docx_apply_page_setup(doc)
    _docx_style_normal(doc)

    logo_raw = _fetch_url_bytes(sender.get("logo_url"))

    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.style = "Table Grid"
    left_cell = tbl.cell(0, 0)
    right_cell = tbl.cell(0, 1)
    sb = _sender_styled_blocks(sender)
    rb = _recipient_styled_blocks(q)
    _docx_fill_cell_styled(
        left_cell,
        sb if sb else [(QUOTE_HEADER_HINT_SENDER, False)],
        logo_bytes=logo_raw,
    )
    _docx_fill_cell_styled(
        right_cell,
        rb if rb else [(QUOTE_HEADER_HINT_RECIPIENT, False)],
        logo_bytes=None,
    )

    doc.add_paragraph("")
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Коммерческое предложение")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)
    qnum = (q.get("number") or "").strip()
    if qnum:
        n = doc.add_paragraph()
        n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = n.add_run(f"№ {qnum}")
        nr.font.name = "Times New Roman"
        nr.font.size = Pt(11)

    doc.add_paragraph("")

    intro = (q.get("intro_text") or "").strip()
    gn = (q.get("greeting_name") or "").strip()
    if intro:
        _para_lines_docx(doc, intro, justify=True)
    elif gn:
        p = doc.add_paragraph(f"Уважаемый {gn}!")
        for rr in p.runs:
            rr.font.name = "Times New Roman"
            rr.font.size = Pt(12)
    else:
        p = doc.add_paragraph("Уважаемый заказчик!")
        for rr in p.runs:
            rr.font.name = "Times New Roman"
            rr.font.size = Pt(12)

    qd_ru = format_quote_date_ru(q.get("quote_date"))
    if qd_ru:
        dp = doc.add_paragraph(qd_ru)
        dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for rr in dp.runs:
            rr.font.name = "Times New Roman"
            rr.font.size = Pt(12)

    doc.add_paragraph("")

    fig = 1
    for ii, it in enumerate(items):
        if it.get("intro"):
            _para_lines_docx(doc, it["intro"], justify=True)
        elif it.get("title"):
            p = doc.add_paragraph()
            rr = p.add_run(str(it["title"]))
            rr.bold = True
            rr.font.name = "Times New Roman"
            rr.font.size = Pt(12)
        if it.get("kit_text"):
            p = doc.add_paragraph("Комплект поставки:")
            for rr in p.runs:
                rr.bold = True
                rr.font.name = "Times New Roman"
                rr.font.size = Pt(12)
            _para_lines_docx(doc, it["kit_text"])
        if it.get("features_text"):
            p = doc.add_paragraph("Основные особенности установки:")
            for rr in p.runs:
                rr.bold = True
                rr.font.name = "Times New Roman"
                rr.font.size = Pt(12)
            _features_numbered_docx(doc, it["features_text"])
        specs = _spec_rows(it.get("specs"))
        _spec_lines_docx(doc, specs)
        photos = it.get("photo_urls") or []
        if isinstance(photos, list) and photos:
            cap = _item_caption(it)
            cap_p = doc.add_paragraph(f"Рис.{fig} {cap}")
            for rr in cap_p.runs:
                rr.font.name = "Times New Roman"
                rr.font.size = Pt(12)
            fig += 1
            first_img = True
            for u in photos:
                if not u:
                    continue
                if first_img:
                    _docx_add_picture_maybe(doc, str(u), max_width_cm=14.5)
                    first_img = False
                else:
                    bp = doc.add_paragraph(str(u))
                    for rr in bp.runs:
                        rr.font.name = "Times New Roman"
                        rr.font.size = Pt(10)
        doc.add_paragraph("")
        if ii < len(items) - 1:
            doc.add_page_break()

    h = doc.add_paragraph()
    hr = h.add_run("Спецификация на поставку товаров")
    hr.bold = True
    hr.font.name = "Times New Roman"
    hr.font.size = Pt(14)

    cur = (q.get("currency") or "RUB").upper()
    rub = cur in ("RUB", "RUR")
    price_h = "Цена, руб. с\nНДС" if rub else "Цена с НДС"
    sum_h = "Стоимость, руб.\nс НДС" if rub else "Стоимость\nс НДС"
    disc_h = "Стоимость со\nскидкой, руб.\nс НДС" if rub else "Стоимость со\nскидкой с НДС"

    show_disc = q.get("show_discount_column", True)
    header = ["№\nп/п", "Наименование\nоборудования", price_h, "Количество", sum_h]
    if show_disc:
        header.append(disc_h)

    st2 = doc.add_table(rows=len(items) + 2, cols=len(header))
    st2.style = "Table Grid"
    for ci, htxt in enumerate(header):
        st2.cell(0, ci).text = htxt
    sum_full = 0.0
    sum_disc = 0.0
    for ri, it in enumerate(items, start=1):
        name = (it.get("title") or "").strip()
        if it.get("model"):
            name = f"{name}\n{it['model']}".strip()
        pw = float(it.get("price_with_vat") or 0)
        qty = float(it.get("quantity") or 1)
        lt = float(it.get("line_total_with_vat") or 0)
        ld = float(it.get("line_total_discounted") or 0)
        sum_full += lt
        sum_disc += ld
        st2.cell(ri, 0).text = str(ri)
        st2.cell(ri, 1).text = name or "—"
        st2.cell(ri, 2).text = money_ru(pw)
        st2.cell(ri, 3).text = str(int(qty)) if float(qty).is_integer() else str(qty)
        st2.cell(ri, 4).text = money_ru(lt)
        if show_disc:
            st2.cell(ri, 5).text = money_ru(ld)

    total_row = len(items) + 1
    ncol = len(header)
    st2.cell(total_row, 0).text = "ИТОГО"
    for i in range(1, ncol):
        st2.cell(total_row, i).text = ""
    if show_disc:
        st2.cell(total_row, 4).text = money_ru(sum_full)
        st2.cell(total_row, 5).text = money_ru(sum_disc)
    else:
        st2.cell(total_row, ncol - 1).text = money_ru(sum_full)
    for ci in range(ncol):
        cell = st2.cell(total_row, ci)
        if cell.text:
            for run in cell.paragraphs[0].runs:
                run.bold = True

    _docx_font_table_cells(st2)

    doc.add_paragraph("")
    _footer_terms_docx(doc, q)

    doc.add_paragraph("")
    p_sig = doc.add_paragraph("С уважением,")
    for rr in p_sig.runs:
        rr.font.name = "Times New Roman"
        rr.font.size = Pt(12)
    sn = (sender.get("signer_name") or "").strip()
    sp = (sender.get("signer_position") or "").strip()
    if sn:
        ps = doc.add_paragraph(sn)
        for rr in ps.runs:
            rr.font.name = "Times New Roman"
            rr.font.size = Pt(12)
    if sp:
        ps2 = doc.add_paragraph(sp)
        for rr in ps2.runs:
            rr.font.name = "Times New Roman"
            rr.font.size = Pt(12)
    ph = (sender.get("phone") or "").strip()
    if ph:
        pt = doc.add_paragraph(f"Тел.: {ph}")
        for rr in pt.runs:
            rr.font.name = "Times New Roman"
            rr.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_setup_body_font() -> str:
    """PDF ближе к Word-КП: Times New Roman (+ Bold), иначе Arial / QUOTE_PDF_FONT_TTF."""
    if "QuoteBody" in pdfmetrics.getRegisteredFontNames():
        return "QuoteBody"
    env_ttf = os.environ.get("QUOTE_PDF_FONT_TTF")
    candidates: list[tuple[str | None, str | None]] = [
        (env_ttf, None),
        (r"C:\Windows\Fonts\times.ttf", r"C:\Windows\Fonts\timesbd.ttf"),
        (r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\arialbd.ttf"),
        (r"C:\Windows\Fonts\calibri.ttf", r"C:\Windows\Fonts\calibrib.ttf"),
        (r"C:\Windows\Fonts\arialuni.ttf", None),
    ]
    for reg, bold in candidates:
        if not reg or not os.path.isfile(reg):
            continue
        try:
            pdfmetrics.registerFont(TTFont("QuoteBody", reg))
            if bold and os.path.isfile(bold):
                pdfmetrics.registerFont(TTFont("QuoteBody-Bold", bold))
                pdfmetrics.registerFontFamily(
                    "QuoteBody",
                    normal="QuoteBody",
                    bold="QuoteBody-Bold",
                    italic="QuoteBody",
                    boldItalic="QuoteBody-Bold",
                )
            else:
                pdfmetrics.registerFontFamily(
                    "QuoteBody",
                    normal="QuoteBody",
                    bold="QuoteBody",
                    italic="QuoteBody",
                    boldItalic="QuoteBody",
                )
            return "QuoteBody"
        except Exception:
            continue
    raise RuntimeError(
        "Не найден TTF-шрифт для PDF с кириллицей. Установите Times/Arial или задайте QUOTE_PDF_FONT_TTF."
    )


def build_quote_pdf(ctx: dict[str, Any]) -> bytes:
    font = _pdf_setup_body_font()
    base = ParagraphStyle(
        name="RU",
        fontName=font,
        fontSize=11,
        leading=14,
        spaceAfter=5,
    )
    h_doc = ParagraphStyle(name="HDOC", parent=base, fontSize=16, leading=19, spaceAfter=8, alignment=1)
    h_sec = ParagraphStyle(name="HSEC", parent=base, fontSize=12, leading=15, spaceBefore=8, spaceAfter=4)
    base_justify = ParagraphStyle(name="RUJ", parent=base, alignment=TA_JUSTIFY)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    story: list[Any] = []

    q = ctx["quote"]
    sender = ctx.get("sender") or {}
    items = ctx.get("items") or []

    logo_raw = _fetch_url_bytes(sender.get("logo_url"))

    hint_style = ParagraphStyle(
        name="HINT",
        parent=base,
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#555555"),
        spaceAfter=2,
    )

    left_lines = _sender_lines(sender)
    left_joined = "\n".join(left_lines).strip()
    if left_joined:
        left_para: Any = Paragraph(_esc_xml("\n".join(left_lines)).replace("\n", "<br/>"), base)
    else:
        left_para = Paragraph(_esc_xml(QUOTE_HEADER_HINT_SENDER), hint_style)

    recv_lines = _recipient_lines(q)
    recv_joined = "\n".join(recv_lines).strip()
    if recv_joined:
        right_para: Any = Paragraph(_esc_xml("\n".join(recv_lines)).replace("\n", "<br/>"), base)
    else:
        right_para = Paragraph(_esc_xml(QUOTE_HEADER_HINT_RECIPIENT), hint_style)

    # Логотип в левой колонке шапки (как в Word), а не отдельной строкой по центру страницы
    left_col_rows: list[list[Any]] = []
    if logo_raw:
        img_logo = _flowable_image_fit(logo_raw, 3.6 * cm, 2.0 * cm)
        if img_logo:
            logo_wrap = Table([[img_logo]], colWidths=[8.2 * cm])
            logo_wrap.setStyle(
                TableStyle(
                    [
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            left_col_rows.append([logo_wrap])
    left_col_rows.append([left_para])
    left_tbl = Table(left_col_rows, colWidths=[8.2 * cm])
    left_tbl.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )

    hdr = Table([[left_tbl, right_para]], colWidths=[8.2 * cm, 8.2 * cm])
    hdr.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LINEABOVE", (0, 0), (-1, -1), 0.25, colors.grey),
                ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.grey),
            ]
        )
    )
    story.append(hdr)
    story.append(Spacer(1, 10))

    story.append(Paragraph(_esc_xml("Коммерческое предложение"), h_doc))
    qnum = (q.get("number") or "").strip()
    if qnum:
        story.append(
            Paragraph(
                _esc_xml(f"№ {qnum}"),
                ParagraphStyle(name="num", parent=base, fontSize=11, alignment=1, spaceAfter=4),
            )
        )
    story.append(Spacer(1, 6))

    intro = (q.get("intro_text") or "").strip()
    gn = (q.get("greeting_name") or "").strip()
    if intro:
        story.append(Paragraph(_esc_xml(intro).replace("\n", "<br/>"), base_justify))
    elif gn:
        story.append(Paragraph(_esc_xml(f"Уважаемый {gn}!"), base))
    else:
        story.append(Paragraph(_esc_xml("Уважаемый заказчик!"), base))

    qd_ru = format_quote_date_ru(q.get("quote_date"))
    if qd_ru:
        story.append(Paragraph(_esc_xml(qd_ru), ParagraphStyle(name="dr", parent=base, alignment=2)))
    story.append(Spacer(1, 8))

    fig = 1
    for ii, it in enumerate(items):
        if it.get("intro"):
            story.append(Paragraph(_esc_xml(str(it["intro"])).replace("\n", "<br/>"), base_justify))
        elif it.get("title"):
            story.append(Paragraph(f"<b>{_esc_xml(str(it['title']))}</b>", base))
        if it.get("kit_text"):
            story.append(Paragraph(_esc_xml("Комплект поставки:"), base))
            story.append(Paragraph(_esc_xml(str(it["kit_text"])).replace("\n", "<br/>"), base))
        if it.get("features_text"):
            story.append(Paragraph(_esc_xml("Основные особенности установки:"), base))
            for i2, ln in enumerate([x.strip() for x in str(it["features_text"]).splitlines() if x.strip()], start=1):
                story.append(Paragraph(_esc_xml(f"{i2}. {ln}"), base))
        specs = _spec_rows(it.get("specs"))
        if specs:
            story.append(Paragraph("<b>" + _esc_xml("ТЕХНИЧЕСКИЕ ХАРАКТЕРИСТИКИ") + "</b>", base))
            for k, v in specs:
                line = f"{k}: {v}" if k and v else (k or v)
                story.append(Paragraph(_esc_xml(line), base))
        photos = it.get("photo_urls") or []
        if isinstance(photos, list) and photos:
            cap = _item_caption(it)
            story.append(Paragraph(_esc_xml(f"Рис.{fig} {cap}"), base))
            fig += 1
            placed_img = False
            for u in photos:
                if not u:
                    continue
                raw_ph = _fetch_url_bytes(str(u))
                if not placed_img and raw_ph:
                    img_ph = _flowable_image_fit(raw_ph, 14.5 * cm, 11 * cm)
                    if img_ph:
                        story.append(Spacer(1, 4))
                        story.append(img_ph)
                        placed_img = True
                        continue
                story.append(Paragraph("• " + _esc_xml(str(u)), base))
        story.append(Spacer(1, 6))
        if ii < len(items) - 1:
            story.append(PageBreak())

    story.append(Paragraph("<b>" + _esc_xml("Спецификация на поставку товаров") + "</b>", h_sec))

    cur = (q.get("currency") or "RUB").upper()
    rub = cur in ("RUB", "RUR")
    price_h = "Цена, руб. с<br/>НДС" if rub else "Цена с<br/>НДС"
    sum_h = "Стоимость, руб.<br/>с НДС" if rub else "Стоимость<br/>с НДС"
    disc_h = "Стоимость со<br/>скидкой, руб.<br/>с НДС" if rub else "Стоимость со<br/>скидкой с НДС"

    show_disc = q.get("show_discount_column", True)
    cols_txt = ["№<br/>п/п", "Наименование<br/>оборудования", price_h, "Количество", sum_h]
    if show_disc:
        cols_txt.append(disc_h)

    sum_full = 0.0
    sum_disc = 0.0
    spec_rows_pdf: list[list[Any]] = [[Paragraph(c, base) for c in cols_txt]]

    for ri, it in enumerate(items, start=1):
        title = (it.get("title") or "").strip()
        model = (it.get("model") or "").strip()
        if model:
            name_xml = f"{_esc_xml(title)}<br/>{_esc_xml(model)}"
        else:
            name_xml = _esc_xml(title)
        pw = float(it.get("price_with_vat") or 0)
        qty = float(it.get("quantity") or 1)
        lt = float(it.get("line_total_with_vat") or 0)
        ld = float(it.get("line_total_discounted") or 0)
        sum_full += lt
        sum_disc += ld
        qty_s = str(int(qty)) if qty % 1 == 0 else str(qty)
        row = [
            Paragraph(_esc_xml(str(ri)), base),
            Paragraph(name_xml or "—", base),
            Paragraph(_esc_xml(money_ru(pw)), base),
            Paragraph(_esc_xml(qty_s), base),
            Paragraph(_esc_xml(money_ru(lt)), base),
        ]
        if show_disc:
            row.append(Paragraph(_esc_xml(money_ru(ld)), base))
        spec_rows_pdf.append(row)

    def _empty_cell():
        return Paragraph("", base)

    if show_disc:
        total_row_pdf = [
            Paragraph(f"<b>{_esc_xml('ИТОГО')}</b>", base),
            _empty_cell(),
            _empty_cell(),
            _empty_cell(),
            Paragraph(f"<b>{_esc_xml(money_ru(sum_full))}</b>", base),
            Paragraph(f"<b>{_esc_xml(money_ru(sum_disc))}</b>", base),
        ]
    else:
        total_row_pdf = [
            Paragraph(f"<b>{_esc_xml('ИТОГО')}</b>", base),
            _empty_cell(),
            _empty_cell(),
            _empty_cell(),
            Paragraph(f"<b>{_esc_xml(money_ru(sum_full))}</b>", base),
        ]
    spec_rows_pdf.append(total_row_pdf)

    cw = [1.1 * cm, 5.6 * cm, 2.4 * cm, 1.6 * cm, 2.6 * cm]
    if show_disc:
        cw.append(2.7 * cm)
    st2 = Table(spec_rows_pdf, colWidths=cw, repeatRows=1)
    tbl_cmds = [
        ("GRID", (0, 0), (-1, -1), 0.35, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8e8e8")),
        ("FONTNAME", (0, 0), (-1, -1), font),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("ALIGN", (2, 0), (-1, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTNAME", (0, -1), (-1, -1), font),
        ("SPAN", (0, -1), (3, -1)),
    ]
    st2.setStyle(TableStyle(tbl_cmds))
    story.append(st2)
    story.append(Spacer(1, 12))

    for label, val in [
        ("", q.get("terms_price_validity")),
        ("Условия поставки", q.get("terms_delivery")),
        ("Срок поставки", q.get("terms_lead_time")),
        ("Условия оплаты", q.get("terms_payment")),
        ("Страна производства", q.get("terms_production_country")),
        ("Гарантия", q.get("terms_warranty")),
    ]:
        if not val:
            continue
        if label:
            story.append(Paragraph(_esc_xml(f"{label}: {val}"), base))
        else:
            story.append(Paragraph(_esc_xml(str(val)), base))

    cur_note = (q.get("terms_currency_note") or "").strip() or _currency_note(q.get("currency"))
    story.append(Paragraph(_esc_xml(cur_note), base))
    addr_note = (q.get("terms_address_note") or "").strip()
    if addr_note:
        story.append(Paragraph(_esc_xml(addr_note), base))

    story.append(Spacer(1, 14))
    story.append(Paragraph(_esc_xml("С уважением,"), base))
    sn = (sender.get("signer_name") or "").strip()
    sp = (sender.get("signer_position") or "").strip()
    if sn:
        story.append(Paragraph(_esc_xml(sn), base))
    if sp:
        story.append(Paragraph(_esc_xml(sp), base))
    ph = (sender.get("phone") or "").strip()
    if ph:
        story.append(Paragraph(_esc_xml(f"Тел.: {ph}"), base))

    doc.build(story)
    return buf.getvalue()


def safe_filename_part(text: str | None, fallback: str) -> str:
    if not text:
        return fallback
    cleaned = re.sub(r"[^\w\-]+", "_", str(text), flags=re.UNICODE)
    return cleaned[:80] or fallback
